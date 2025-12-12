#!/usr/bin/env python3
"""
A script to extract all content (text and tables) from a PDF file and save it
into a structured Microsoft Word (.docx) document.

This script uses the Docling library to parse the PDF and preserve its structure,
including headings, paragraphs, lists, and tables.

Usage:
    python extract_full_document_to_word.py <path_to_pdf_file>

Example:
    python extract_full_document_to_word.py my_document.pdf
"""

import logging
import os
import sys
from pathlib import Path

import pandas as pd
from docling.document_converter import DocumentConverter
from docx import Document
from docx.shared import Mm, Pt
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml import OxmlElement
from docx.oxml.ns import qn

_log = logging.getLogger(__name__)

def set_cell_margins(cell, top=0, start=0, bottom=0, end=0):
    """
    Sets the margins (padding) of a table cell to zero.
    """
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    tcMar = OxmlElement('w:tcMar')
    
    for tag, value in [('w:top', top), ('w:start', start), ('w:bottom', bottom), ('w:end', end)]:
        node = OxmlElement(tag)
        node.set(qn('w:w'), str(value))
        node.set(qn('w:type'), 'dxa')
        tcMar.append(node)
    
    tcPr.append(tcMar)

def format_paragraph(paragraph):
    """
    Sets font to 10pt Arial and removes extra spacing.
    """
    paragraph.paragraph_format.space_before = Pt(0)
    paragraph.paragraph_format.space_after = Pt(0)
    paragraph.paragraph_format.line_spacing = 1.0
    
    for run in paragraph.runs:
        run.font.name = 'Arial'
        run.font.size = Pt(10)

def extract_full_document_to_word(pdf_path: str):
    """
    Extracts all content from a PDF and saves it to a Word document.

    Args:
        pdf_path: The path to the PDF file.
    """
    logging.basicConfig(level=logging.INFO)

    if "TESSDATA_PREFIX" not in os.environ:
        tessdata_prefix = "/usr/share/tesseract-ocr/5/tessdata/"
        if os.path.exists(tessdata_prefix):
            os.environ["TESSDATA_PREFIX"] = tessdata_prefix

    input_doc_path = Path(pdf_path)
    if not input_doc_path.exists():
        _log.error(f"Error: PDF file not found at '{input_doc_path}'")
        return

    output_docx_path = Path(f"{input_doc_path.stem}_full_content.docx")

    _log.info(f"Processing PDF: {input_doc_path}")
    _log.info(f"Output will be saved to: {output_docx_path}")

    doc_converter = DocumentConverter()
    conv_res = doc_converter.convert(input_doc_path)
    docling_doc = conv_res.document

    # Create a mapping of self_ref to element for easy lookup
    element_map = {}
    for element_type in ["texts", "tables", "groups", "pictures"]:
        for element in getattr(docling_doc, element_type, []):
            element_map[element.self_ref] = element

    document = Document()
    
    # Set Page Margins (5mm)
    for section in document.sections:
        section.top_margin = Mm(5)
        section.bottom_margin = Mm(5)
        section.left_margin = Mm(5)
        section.right_margin = Mm(5)

    # Process elements in the order they appear in the document body
    for child_ref in docling_doc.body.children:
        element_cref = child_ref.cref
        element_type = element_cref.split('/')[1] # e.g., #/texts/0 -> texts
        element = element_map.get(element_cref)
        if not element:
            continue

        if element_type == "texts":
            p = document.add_paragraph(element.text)
            format_paragraph(p)
            
            if hasattr(element, 'label') and element.label == "section_header":
                for run in p.runs:
                    run.bold = True
            elif hasattr(element, 'label') and element.label == "list_item":
                p.style = 'List Bullet'
                format_paragraph(p) # Re-apply format to override style defaults if needed

        elif element_type == "tables":
            table_df: pd.DataFrame = element.export_to_dataframe(doc=docling_doc)
            
            if not table_df.empty:
                # Add a table to the document
                rows = len(table_df) + 1 # +1 for header
                cols = len(table_df.columns)
                doc_table = document.add_table(rows=rows, cols=cols)
                doc_table.style = 'Table Grid'
                doc_table.autofit = True
                doc_table.alignment = WD_TABLE_ALIGNMENT.CENTER

                # Header
                for j, col_name in enumerate(table_df.columns):
                    cell = doc_table.cell(0, j)
                    cell.text = str(col_name)
                    set_cell_margins(cell, 0, 0, 0, 0)
                    for p in cell.paragraphs:
                        format_paragraph(p)
                        for run in p.runs:
                            run.bold = True

                # Data
                for i, row in enumerate(table_df.itertuples(index=False)):
                    for j, cell_value in enumerate(row):
                        cell = doc_table.cell(i+1, j)
                        cell.text = str(cell_value) if cell_value is not None else ""
                        set_cell_margins(cell, 0, 0, 0, 0)
                        for p in cell.paragraphs:
                            format_paragraph(p)
                
                document.add_paragraph() # Spacer

    document.save(output_docx_path)
    _log.info(f"Successfully extracted full document content to '{output_docx_path}'.")
    return output_docx_path

if __name__ == "__main__":
    if len(sys.argv) != 2:
        print("Usage: python extract_full_document_to_word.py <path_to_pdf_file>")
        print("Example: python extract_full_document_to_word.py my_document.pdf")
        sys.exit(1)

    pdf_file = sys.argv[1]
    extract_full_document_to_word(pdf_file)
