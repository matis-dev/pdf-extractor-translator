#!/usr/bin/env python3
"""
A script to extract tables from a PDF file and save them into a Microsoft Word (.docx) document.

This script uses the Docling library to parse the PDF and extract tables, and then
uses the python-docx library to create a Word document containing the extracted tables.

Usage:
    python extract_tables_to_word.py <path_to_pdf_file>

Example:
    python extract_tables_to_word.py my_document.pdf
"""

import logging
import os
import sys
from pathlib import Path

import pandas as pd
from docling.document_converter import DocumentConverter
from docx import Document

_log = logging.getLogger(__name__)


def extract_tables_to_word(pdf_path: str):
    """
    Extracts tables from a PDF file and saves them into a Word document.

    Args:
        pdf_path: The path to the PDF file.
    """
    logging.basicConfig(level=logging.INFO)

    # Set TESSDATA_PREFIX if not already set (important for tesserocr)
    if "TESSDATA_PREFIX" not in os.environ:
        tessdata_prefix = "/usr/share/tesseract-ocr/5/tessdata/"
        if os.path.exists(tessdata_prefix):
            os.environ["TESSDATA_PREFIX"] = tessdata_prefix
        else:
            _log.warning(
                f"TESSDATA_PREFIX not set and default path '{tessdata_prefix}' not found. "
                "OCR may fail if Tesseract is not configured correctly."
            )

    input_doc_path = Path(pdf_path)
    if not input_doc_path.exists():
        _log.error(f"Error: PDF file not found at '{input_doc_path}'")
        return

    output_docx_path = Path(f"{input_doc_path.stem}_tables.docx")

    _log.info(f"Processing PDF: {input_doc_path}")
    _log.info(f"Output will be saved to: {output_docx_path}")

    doc_converter = DocumentConverter()
    conv_res = doc_converter.convert(input_doc_path)

    if not conv_res.document.tables:
        _log.info("No tables found in the document.")
        return

    # Create a new Word document
    document = Document()
    document.add_heading(f"Tables Extracted from {input_doc_path.name}", 0)

    for table_ix, table in enumerate(conv_res.document.tables):
        document.add_heading(f"Table {table_ix + 1}", level=1)
        table_df: pd.DataFrame = table.export_to_dataframe(doc=conv_res.document)

        # Add a table to the document
        doc_table = document.add_table(rows=1, cols=len(table_df.columns))
        doc_table.style = 'Table Grid'

        # Add the header row
        for j, col_name in enumerate(table_df.columns):
            doc_table.cell(0, j).text = str(col_name)

        # Add the data rows
        for i, row in table_df.iterrows():
            row_cells = doc_table.add_row().cells
            for j, cell_value in enumerate(row):
                row_cells[j].text = str(cell_value)
    
    # Save the document
    document.save(output_docx_path)
    _log.info(f"Successfully extracted {len(conv_res.document.tables)} tables to '{output_docx_path}'.")


if __name__ == "__main__":
    if len(sys.argv) != 2:
        print("Usage: python extract_tables_to_word.py <path_to_pdf_file>")
        print("Example: python extract_tables_to_word.py my_document.pdf")
        sys.exit(1)

    pdf_file = sys.argv[1]
    extract_tables_to_word(pdf_file)
