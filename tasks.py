from celery import shared_task
import os
import shutil
from pathlib import Path
from docx import Document
from extract_full_document_to_word import extract_full_document_to_word
from extract_tables_to_csv import extract_tables as extract_tables_to_csv
from translation_utils import translate_text, install_languages
import subprocess

def get_unique_filename(directory, filename):
    """Generates a unique filename by appending a counter if the file exists.

    Args:
        directory (str): The directory specifically where the file will be saved.
        filename (str): The original desired filename.

    Returns:
        str: A unique filename that does not conflict with existing files in the directory.
    """
    base, ext = os.path.splitext(filename)
    counter = 1
    new_filename = filename
    while os.path.exists(os.path.join(directory, new_filename)):
        new_filename = f"{base}({counter}){ext}"
        counter += 1
    return new_filename

# Ensure languages are installed (could be done on startup)
try:
    install_languages()
except Exception as e:
    print(f"Warning: Could not install languages: {e}")

@shared_task(bind=True)
def process_pdf_task(self, extraction_type, filename, upload_folder, output_folder, target_lang=None, source_lang='en'):
    """Celery background task to handle PDF extraction and optional translation.

    Args:
        extraction_type (str): The mode of extraction ('word', 'odt', 'csv').
        filename (str): Name of the uploaded PDF file.
        upload_folder (str): Path to the folder containing uploads.
        output_folder (str): Path to the folder where results should be saved.
        target_lang (str, optional): ISO code for translation (e.g., 'es'). Defaults to None.
        source_lang (str, optional): ISO code of original document. Defaults to 'en'.

    Returns:
        dict: A result dictionary containing status and result_file path.
    """
    def update_progress(state, meta):
        self.update_state(state=state, meta=meta)

    return run_pdf_extraction(
        extraction_type, 
        filename, 
        upload_folder, 
        output_folder, 
        target_lang, 
        source_lang, 
        progress_callback=update_progress
    )

def run_pdf_extraction(extraction_type, filename, upload_folder, output_folder, target_lang=None, source_lang='en', progress_callback=None):
    """Core PDF extraction logic, decoupled from Celery for testability.

    This function handles the heavy lifting of calling Docling, processing tables,
    and performing text translation on the resulting document.

    Args:
        extraction_type (str): Type of output requested.
        filename (str): Input filename.
        upload_folder (str): Source directory.
        output_folder (str): Destination directory.
        target_lang (str, optional): Target language for translation.
        source_lang (str): Source language.
        progress_callback (callable, optional): function(state, meta) to report progress.

    Returns:
        dict: Completion status and final filename.
    """
    if progress_callback:
        progress_callback('PROCESSING', {'status': 'Starting extraction...', 'current': 0, 'total': 100})
    
    pdf_path = os.path.join(upload_folder, filename)
    result_file = None

    try:
        if extraction_type == 'word' or extraction_type == 'odt':
            if progress_callback:
                progress_callback('PROCESSING', {'status': 'Extracting content...', 'current': 10, 'total': 100})
            
            output_path = extract_full_document_to_word(pdf_path)
            
            if progress_callback:
                progress_callback('PROCESSING', {'status': 'Extraction complete. Preparing translation...', 'current': 30, 'total': 100})
            
            # Translate if requested
            if target_lang and target_lang != 'none':
                 if progress_callback:
                    progress_callback('PROCESSING', {'status': f'Translating to {target_lang}...', 'current': 30, 'total': 100})
                 
                 doc = Document(output_path)
                 
                 # Count total items to translate for progress
                 total_items = len(doc.paragraphs) + sum(len(row.cells) * len(cell.paragraphs) for table in doc.tables for row in table.rows for cell in row.cells)
                 processed_items = 0
                 
                 # Translate paragraphs
                 for para in doc.paragraphs:
                     if para.text.strip():
                         para.text = translate_text(para.text, target_lang, source_lang)
                     processed_items += 1
                     if processed_items % 10 == 0 and progress_callback:
                         progress = 30 + int((processed_items / total_items) * 60) # 30% to 90%
                         progress_callback('PROCESSING', {'status': f'Translating... ({int(processed_items/total_items*100)}%)', 'current': progress, 'total': 100})

                 # Translate tables
                 for table in doc.tables:
                     for row in table.rows:
                         for cell in row.cells:
                             for para in cell.paragraphs:
                                 if para.text.strip():
                                     para.text = translate_text(para.text, target_lang, source_lang)
                                 processed_items += 1
                                 if processed_items % 10 == 0 and progress_callback:
                                     progress = 30 + int((processed_items / total_items) * 60)
                                     progress_callback('PROCESSING', {'status': f'Translating tables... ({int(processed_items/total_items*100)}%)', 'current': progress, 'total': 100})
                 
                 doc.save(output_path)
                 
                 # Append language code to filename
                 directory, filename = os.path.split(output_path)
                 name, ext = os.path.splitext(filename)
                 new_filename = f"{name}_{target_lang}{ext}"
                 new_output_path = os.path.join(directory, new_filename)
                 os.rename(output_path, new_output_path)
                 output_path = new_output_path

            if extraction_type == 'odt':
                if progress_callback:
                    progress_callback('PROCESSING', {'status': 'Converting to ODT...'})
                odt_path = output_path.replace('.docx', '.odt')
                subprocess.run(['pandoc', output_path, '-o', odt_path], check=True)
                os.remove(output_path) # Remove intermediate DOCX
                output_path = odt_path

            # Move the file to the output folder
            final_filename = get_unique_filename(output_folder, os.path.basename(output_path))
            destination = os.path.join(output_folder, final_filename)
            
            shutil.move(output_path, destination)
            result_file = final_filename

        elif extraction_type == 'csv':
            if progress_callback:
                progress_callback('PROCESSING', {'status': 'Extracting tables...', 'current': 10, 'total': 100})
            output_dir_path = extract_tables_to_csv(pdf_path)
            
            # Zip the contents of the output directory
            zip_filename = f"{Path(pdf_path).stem}_csv_files"
            shutil.make_archive(zip_filename, 'zip', output_dir_path)
            
            # Move the zip file to the output folder
            final_filename = get_unique_filename(output_folder, f"{zip_filename}.zip")
            destination = os.path.join(output_folder, final_filename)
            
            shutil.move(f"{zip_filename}.zip", destination)
            result_file = final_filename

            # Clean up the original CSV directory
            shutil.rmtree(output_dir_path)
        
        return {'status': 'Completed', 'result_file': result_file}
    except Exception as e:
        # In a real app, you might want to log this better
        if progress_callback:
            progress_callback('FAILURE', {'status': 'Failed', 'error': str(e)})
        return {'status': 'Failed', 'error': str(e)}
